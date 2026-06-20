#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
旅行计划 Word 文档生成器

读取结构化行程 JSON，生成排版美观、可编辑的 Word 行程手册。
使用 python-docx 库，支持标题、表格、颜色、样式。
"""

import json
import argparse
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ============================================================
# 配色方案
# ============================================================
COLOR_PRIMARY = RGBColor(0x2E, 0x86, 0xAB)      # 海洋蓝
COLOR_PRIMARY_DARK = RGBColor(0x1B, 0x5E, 0x7B) # 深蓝
COLOR_SECONDARY = RGBColor(0xF1, 0x8F, 0x01)    # 暖橙
COLOR_ACCENT = RGBColor(0xC7, 0x3E, 0x1D)       # 红色
COLOR_SUCCESS = RGBColor(0x27, 0xAE, 0x60)      # 绿色
COLOR_TEXT = RGBColor(0x2D, 0x34, 0x36)         # 正文
COLOR_TEXT_LIGHT = RGBColor(0x63, 0x6E, 0x72)   # 浅文字
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# 活动分类颜色
CATEGORY_COLORS = {
    'transport': COLOR_PRIMARY,
    'sight': COLOR_SECONDARY,
    'food': COLOR_ACCENT,
    'hotel': COLOR_SUCCESS,
    'shopping': RGBColor(0x8E, 0x44, 0xAD),
    'activity': RGBColor(0xE6, 0x7E, 0x22),
}

CATEGORY_NAMES = {
    'transport': '交通',
    'sight': '景点',
    'food': '餐饮',
    'hotel': '住宿',
    'shopping': '购物',
    'activity': '活动',
}


# ============================================================
# 工具函数
# ============================================================
def set_cell_shading(cell, color_hex):
    """设置表格单元格背景色。"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_run(paragraph, text, bold=False, color=None, size=None, italic=False):
    """向段落中添加格式化文本片段。"""
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.italic = italic
    return run


def create_styled_table(doc, headers, rows, col_widths=None):
    """创建带样式的表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = COLOR_WHITE
        run.font.size = Pt(12)
        set_cell_shading(cell, '2E86AB')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(12)
            # 交替行背景
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F0F4F8')

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


# ============================================================
# 文档构建主逻辑
# ============================================================
def build_docx(data, output_path):
    """根据结构化数据生成 Word 行程手册。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    font.color.rgb = COLOR_TEXT
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置标题样式（四号 14pt）
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.color.rgb = COLOR_PRIMARY_DARK
        heading_style.font.size = Pt(14)
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============ 封面 ============
    # 添加空行推到中间位置
    for _ in range(6):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(data.get('title', '旅行计划'))
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = COLOR_PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(data.get('subtitle', '— 你的旅程，由 AI 规划 —'))
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_TEXT_LIGHT

    doc.add_paragraph('')

    # 基本信息
    info_items = []
    if data.get('date_range'):
        info_items.append(data['date_range'])
    if data.get('travelers'):
        t = data['travelers']
        if isinstance(t, dict):
            parts = [f"{t.get('count', '?')}人出行"]
            if t.get('type'):
                parts.append(t['type'])
            if t.get('discounts'):
                parts.append('优惠: ' + '、'.join(t['discounts']))
            info_items.append(' · '.join(parts))
        else:
            info_items.append(str(t))
    if data.get('destinations'):
        info_items.append(' · '.join(data['destinations']))

    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_TEXT_LIGHT

    doc.add_page_break()

    # ============ 行程概览 ============
    doc.add_heading('行程概览', level=1)

    days = data.get('days', [])
    for i, day in enumerate(days):
        day_num = i + 1
        date_str = day.get('date', '')
        theme = day.get('theme', '')
        city = day.get('city') or (data.get('destinations') or [''])[0]
        city_tag = f' [{city}]' if city else ''

        p = doc.add_paragraph()
        add_formatted_run(p, f'Day {day_num}  ', bold=True, color=COLOR_PRIMARY, size=12)
        add_formatted_run(p, f'{date_str}{city_tag}', bold=True, size=12)

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        add_formatted_run(p2, theme, color=COLOR_TEXT_LIGHT, size=12)

    # 天气摘要
    weather = data.get('weather') or {}
    if weather:
        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        high, low = weather.get('high', ''), weather.get('low', '')
        rain = weather.get('rain', '')
        dressing = weather.get('dressing', '')
        risk = weather.get('risk', '')
        parts = [f'{weather.get("period", "")}']
        if high and low:
            parts.append(f'高温 {high}°C / 低温 {low}°C')
        if rain and rain != '无':
            parts.append(rain)
        if dressing:
            parts.append(f'穿: {dressing}')
        add_formatted_run(p, ' · '.join(parts), color=COLOR_TEXT_LIGHT, size=12)
        if risk:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p2, f'⚠ {risk}', color=COLOR_SECONDARY, size=12)

    # 预算预览
    budget = data.get('budget') or {}
    grand_total = budget.get('grand_total', 0)
    if grand_total:
        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_formatted_run(p, f'预估总费用: ', bold=True, size=13)
        add_formatted_run(p, f'¥{grand_total:,}', bold=True, color=COLOR_ACCENT, size=16)

    doc.add_page_break()

    # ============ 每日行程 ============
    doc.add_heading('每日行程', level=1)

    for i, day in enumerate(days):
        day_num = i + 1
        date_str = day.get('date', '')
        theme = day.get('theme', '')
        city = day.get('city') or (data.get('destinations') or [''])[0]

        # Day 标题
        heading = doc.add_heading(f'Day {day_num}  {date_str}', level=2)
        for run in heading.runs:
            run.font.color.rgb = COLOR_PRIMARY

        # 主题
        p = doc.add_paragraph()
        add_formatted_run(p, theme, italic=True, color=COLOR_TEXT_LIGHT, size=12)

        # 住宿信息
        accom = day.get('accommodation') or {}
        if accom and accom.get('name'):
            p = doc.add_paragraph()
            add_formatted_run(p, '住宿: ', bold=True, color=COLOR_SUCCESS, size=12)
            hotel_text = accom['name']
            if accom.get('address'):
                hotel_text += f' | {accom["address"]}'
            if accom.get('price'):
                hotel_text += f' | {accom["price"]}'
            if accom.get('room_plan'):
                hotel_text += f' | {accom["room_plan"]}'
            add_formatted_run(p, hotel_text, color=COLOR_SUCCESS, size=12)

        # 活动按时间段分组
        activities = day.get('activities', [])
        time_periods = {'morning': [], 'afternoon': [], 'evening': []}

        for idx, act in enumerate(activities):
            time_str = act.get('time', '')
            hour = None
            if time_str:
                try:
                    parts = re.split(r'[:\uff1a]', time_str)
                    hour = int(parts[0])
                except (ValueError, IndexError):
                    pass

            if hour is not None:
                if hour < 12:
                    period = 'morning'
                elif hour < 18:
                    period = 'afternoon'
                else:
                    period = 'evening'
            else:
                # 无 time 字段时按数组位置分配：前1/3上午，中1/3下午，后1/3晚上
                n = len(activities)
                ratio = idx / max(n - 1, 1)
                if ratio < 0.34:
                    period = 'morning'
                elif ratio < 0.67:
                    period = 'afternoon'
                else:
                    period = 'evening'
            time_periods[period].append(act)

        period_names = {
            'morning': '上午',
            'afternoon': '下午',
            'evening': '晚上',
        }

        for period_key in ['morning', 'afternoon', 'evening']:
            period_acts = time_periods[period_key]
            if not period_acts:
                continue

            # 时间段标题
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            add_formatted_run(p, f'  {period_names[period_key]}', bold=True, color=COLOR_PRIMARY_DARK, size=12)

            for act in period_acts:
                time_str = act.get('time', '')
                title = act.get('title', '')
                detail = act.get('detail', '')
                cost = act.get('cost', '')
                category = act.get('category', 'default')
                duration = act.get('duration', '')
                rating = act.get('rating', '')

                cat_color = CATEGORY_COLORS.get(category, COLOR_TEXT_LIGHT)

                # 交通连接线（从上一个活动到当前活动）
                transit = act.get('transit')
                if transit and (transit.get('options') or []):
                    p_transit = doc.add_paragraph()
                    p_transit.paragraph_format.left_indent = Cm(1.5)
                    p_transit.paragraph_format.space_before = Pt(6)
                    p_transit.paragraph_format.space_after = Pt(2)

                    option_strs = []
                    for opt in transit['options']:
                        parts = [opt.get('mode', '')]
                        if opt.get('duration'):
                            parts.append(opt['duration'])
                        if opt.get('cost'):
                            parts.append(opt['cost'])
                        s = ' '.join(parts)
                        if opt.get('recommended'):
                            s += ' \u2b50'
                        option_strs.append(s)
                    add_formatted_run(p_transit, '\u2192 ' + ' | '.join(option_strs),
                                      color=COLOR_PRIMARY, size=11)

                    rec = transit.get('recommendation', '')
                    if rec:
                        p_rec = doc.add_paragraph()
                        p_rec.paragraph_format.left_indent = Cm(2.0)
                        p_rec.paragraph_format.space_before = Pt(0)
                        p_rec.paragraph_format.space_after = Pt(2)
                        add_formatted_run(p_rec, rec, color=COLOR_TEXT_LIGHT, size=10, italic=True)

                # 活动标题行
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)

                # 分类标签
                cat_name = CATEGORY_NAMES.get(category, '')
                if cat_name:
                    add_formatted_run(p, f'[{cat_name}] ', bold=True, color=cat_color, size=12)

                add_formatted_run(p, title, bold=True, size=12)

                # 评分和时长
                extras = []
                if rating:
                    extras.append(f'推荐指数: {rating}')
                if duration:
                    extras.append(f'建议游玩: {duration}')
                if cost:
                    extras.append(f'费用: {cost}')

                if extras:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Cm(1.5)
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(1)
                    add_formatted_run(p2, ' | '.join(extras), color=COLOR_SECONDARY, size=12)

                # 详情
                if detail:
                    p3 = doc.add_paragraph()
                    p3.paragraph_format.left_indent = Cm(1.5)
                    p3.paragraph_format.space_before = Pt(0)
                    p3.paragraph_format.space_after = Pt(2)
                    add_formatted_run(p3, detail, color=COLOR_TEXT_LIGHT, size=12)

                # 景区内游览路线（大型景区 route 字段）
                route = act.get('route') or {}
                if route and category == 'sight':
                    route_table = doc.add_table(rows=1, cols=1)
                    route_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    route_cell = route_table.rows[0].cells[0]
                    route_cell.text = ''
                    set_cell_shading(route_cell, 'EDF6F9')

                    p_head = route_cell.paragraphs[0]
                    p_head.paragraph_format.space_after = Pt(2)
                    run_head = p_head.add_run('🗺 景区内游览路线')
                    run_head.bold = True
                    run_head.font.size = Pt(12)
                    run_head.font.color.rgb = COLOR_PRIMARY

                    must_see = route.get('must_see') or []
                    if must_see:
                        p_see = route_cell.add_paragraph()
                        p_see.paragraph_format.space_after = Pt(1)
                        run_l = p_see.add_run('必看: ')
                        run_l.bold = True
                        run_l.font.size = Pt(11)
                        run_l.font.color.rgb = COLOR_TEXT
                        run_v = p_see.add_run(' · '.join(must_see))
                        run_v.font.size = Pt(11)
                        run_v.font.color.rgb = COLOR_TEXT

                    path = route.get('path', '')
                    if path:
                        p_path = route_cell.add_paragraph()
                        p_path.paragraph_format.space_after = Pt(1)
                        run_l = p_path.add_run('路线: ')
                        run_l.bold = True
                        run_l.font.size = Pt(11)
                        run_l.font.color.rgb = COLOR_TEXT
                        run_v = p_path.add_run(path)
                        run_v.font.size = Pt(11)
                        run_v.font.color.rgb = COLOR_TEXT_LIGHT

                    skip = route.get('skip', '')
                    if skip:
                        p_skip = route_cell.add_paragraph()
                        p_skip.paragraph_format.space_after = Pt(0)
                        run_l = p_skip.add_run('可跳过: ')
                        run_l.bold = True
                        run_l.font.size = Pt(11)
                        run_l.font.color.rgb = COLOR_TEXT
                        run_v = p_skip.add_run(skip)
                        run_v.font.size = Pt(11)
                        run_v.font.color.rgb = COLOR_TEXT_LIGHT

        # Day 分隔线
        if i < len(days) - 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p, '─' * 40, color=RGBColor(0xB2, 0xBE, 0xC3), size=9)

    doc.add_page_break()

    # ============ 预算明细 ============
    doc.add_heading('预算明细', level=1)

    if budget and budget.get('grand_total'):
        category_names = {
            'transport': '交通',
            'accommodation': '住宿',
            'tickets': '门票',
            'food': '餐饮',
            'other': '其他',
        }

        headers = ['费用类别', '明细', '金额']
        rows = []
        for key, name in category_names.items():
            if key in budget:
                cat = budget[key]
                total = cat.get('total', 0)
                items = cat.get('items', [])
                detail = '、'.join(items[:4])
                if len(items) > 4:
                    detail += '...'
                rows.append([name, detail, f'¥{total:,}'])

        # 总计行
        rows.append(['', '', ''])  # 空行
        rows.append(['总计', '', f'¥{budget["grand_total"]:,}'])

        table = create_styled_table(doc, headers, rows, col_widths=[3, 8, 3])

        # 高亮总计行
        last_row = table.rows[-1]
        for cell in last_row.cells:
            set_cell_shading(cell, 'FFF0E0')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = COLOR_ACCENT

        # 人均费用
        if grand_total and data.get('travelers'):
            t = data['travelers']
            if isinstance(t, dict):
                people = t.get('count', 0)
            else:
                nums = re.findall(r'\d+', str(t))
                people = int(nums[0]) if nums else 0
            if people > 0:
                per_person = round(grand_total / people)
                doc.add_paragraph('')
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_run(p, f'人均费用: ¥{per_person:,}（{people}人出行）',
                                  bold=True, color=COLOR_TEXT, size=12)

    doc.add_page_break()

    # ============ 实用贴士 ============
    doc.add_heading('实用贴士', level=1)
    tips = data.get('tips') or {}

    # 行李清单
    packing = tips.get('packing', [])
    if packing:
        doc.add_heading('行李清单', level=2)
        # 3列表格
        items_per_row = 3
        for i in range(0, len(packing), items_per_row):
            chunk = packing[i:i + items_per_row]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            for j, item in enumerate(chunk):
                prefix = '  ' if j > 0 else ''
                add_formatted_run(p, f'{prefix}  {item}    ', size=12)

    # 注意事项
    notes = tips.get('notes', [])
    if notes:
        doc.add_heading('注意事项', level=2)
        for note in notes:
            p = doc.add_paragraph()
            add_formatted_run(p, f'  {note}', size=12)

    # 紧急联系
    emergency = tips.get('emergency', [])
    if emergency:
        doc.add_heading('紧急联系', level=2)
        for item in emergency:
            p = doc.add_paragraph()
            add_formatted_run(p, f'  {item}', size=12)

    # ============ 美食备选 ============
    extra_restaurants = data.get('extra_restaurants') or []
    if extra_restaurants:
        doc.add_page_break()
        doc.add_heading('更多美食选择', level=1)

        p = doc.add_paragraph()
        add_formatted_run(p, '以下餐厅供你灵活选择，行程中安排的餐厅是当天交通最方便的，这些是其他好评餐厅，想去哪家随时可以换。',
                          color=COLOR_TEXT_LIGHT, size=12)
        doc.add_paragraph('')

        headers = ['餐厅', '菜系', '人均', '地址', '推荐理由']
        rows = []
        for r in extra_restaurants:
            rows.append([
                r.get('name', ''),
                r.get('cuisine', ''),
                r.get('price', ''),
                r.get('address', ''),
                r.get('recommendation', ''),
            ])
        create_styled_table(doc, headers, rows, col_widths=[3.5, 2.5, 2, 3.5, 4.5])

    # ============ 尾页 ============
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, '祝旅途愉快!', bold=True, color=COLOR_PRIMARY, size=24)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer = data.get('disclaimer', '本行程由 AI 旅行规划助手生成，实际出行前请核实门票价格、交通班次等信息。')
    add_formatted_run(p, disclaimer, color=COLOR_TEXT_LIGHT, size=10)

    # 保存
    doc.save(output_path)
    return output_path


# ============================================================
# 数据校验
# ============================================================
def validate_data(data):
    """校验行程 JSON 数据，返回错误列表。"""
    errors = []

    if not isinstance(data, dict):
        return ['JSON 顶层必须是对象（dict）']

    if not data.get('title'):
        errors.append('缺少必填字段: title')
    if not data.get('days'):
        errors.append('缺少必填字段: days')
    elif not isinstance(data['days'], list):
        errors.append('days 必须是数组')
    else:
        valid_categories = set(CATEGORY_COLORS.keys())
        for i, day in enumerate(data['days']):
            prefix = f'days[{i}]'
            if not day.get('date'):
                errors.append(f'{prefix} 缺少 date 字段')
            activities = day.get('activities', [])
            for j, act in enumerate(activities):
                act_prefix = f'{prefix}.activities[{j}]'
                if not act.get('title'):
                    errors.append(f'{act_prefix} 缺少 title 字段')
                # category 枚举检查
                cat = act.get('category', '')
                if cat and cat not in valid_categories:
                    errors.append(f'{act_prefix} category 值 "{cat}" 非法，合法值: {", ".join(sorted(valid_categories))}')
                # time 格式检查（可选字段）
                time_val = act.get('time', '')
                if time_val and isinstance(time_val, str) and not re.fullmatch(r'[0-9]{1,2}:[0-9]{2}', time_val):
                    errors.append(f'{act_prefix} time 格式错误 "{time_val}"，应为 HH:MM')
                # cost 格式检查
                cost = act.get('cost', '')
                if cost and isinstance(cost, str) and '¥' not in cost and cost not in ('免费', 'free', '0'):
                    errors.append(f'{act_prefix} cost "{cost}" 缺少 ¥ 符号，也不是"免费"')
                elif cost and not isinstance(cost, str):
                    errors.append(f'{act_prefix} cost 必须是字符串，当前类型为 {type(cost).__name__}')
                # rating 格式检查
                rating = act.get('rating', '')
                if rating and isinstance(rating, str) and not re.fullmatch(r'[★☆]+', rating):
                    errors.append(f'{act_prefix} rating "{rating}" 格式错误，应仅含 ★☆')
                elif rating and not isinstance(rating, str):
                    errors.append(f'{act_prefix} rating 必须是字符串，当前类型为 {type(rating).__name__}')
                # detail 长度警告
                detail = act.get('detail') or ''
                if detail and len(detail) > 100:
                    errors.append(f'{act_prefix} detail 过长（{len(detail)}字），建议 100 字以内')
                # transit 校验（非首个活动必须有 transit）
                if j > 0:
                    transit = act.get('transit')
                    if not transit:
                        errors.append(f'{act_prefix} 缺少 transit 字段（非首个活动必须标注交通方式）')
                    else:
                        opts = transit.get('options') or []
                        # transport/hotel 类活动至少 1 种方案即可，其他类型至少 2 种
                        act_cat = act.get('category', '')
                        min_opts = 1 if act_cat in ('transport', 'hotel') else 2
                        if len(opts) < min_opts:
                            errors.append(f'{act_prefix} transit.options 至少需要 {min_opts} 种交通方式，当前 {len(opts)} 种')
                        rec_count = sum(1 for o in opts if o.get('recommended'))
                        if rec_count > 1:
                            errors.append(f'{act_prefix} transit 只能有 1 个推荐方案，当前有 {rec_count} 个')

        # 预算校验
        budget = data.get('budget') or {}
        if budget:
            category_keys = {'transport', 'accommodation', 'tickets', 'food', 'other'}
            grand_total = budget.get('grand_total', 0)
            computed_total = sum(
                budget.get(k, {}).get('total', 0)
                for k in category_keys if k in budget
            )
            if abs(grand_total - computed_total) > 1:
                errors.append(
                    f'budget.grand_total ({grand_total}) 与分类合计 ({computed_total}) 不一致'
                )

    # tips 校验
    tips = data.get('tips') or {}
    packing = tips.get('packing') or []
    if packing and not (9 <= len(packing) <= 15):
        errors.append(f'tips.packing 数量 {len(packing)} 不在建议范围（9-15）')

    # weather 校验
    weather = data.get('weather')
    if weather and isinstance(weather, dict):
        for key in ('high', 'low'):
            val = weather.get(key)
            if val is not None and not isinstance(val, (int, float)):
                errors.append(f'weather.{key} 应为数字，当前为 "{val}"')

    # extra_restaurants 校验
    extra_r = data.get('extra_restaurants') or []
    if extra_r:
        if not (5 <= len(extra_r) <= 10):
            errors.append(f'extra_restaurants 数量 {len(extra_r)} 不在建议范围（5-10）')
        for i, r in enumerate(extra_r):
            for field in ('name', 'cuisine', 'price', 'address', 'recommendation'):
                if not r.get(field):
                    errors.append(f'extra_restaurants[{i}] 缺少 {field} 字段')

    return errors


# ============================================================
# 入口
# ============================================================
def main():
    parser = argparse.ArgumentParser(description='旅行计划 Word 文档生成器')
    parser.add_argument('--input', required=True, help='输入 JSON 文件（行程数据）')
    parser.add_argument('--output', default='trip_plan.docx', help='输出 Word 文件路径')
    parser.add_argument('--force', action='store_true', help='跳过数据校验，强制生成（调试用）')
    args = parser.parse_args()

    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f'❌ 文件不存在: {args.input}')
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f'❌ JSON 解析失败: {e}')
        sys.exit(1)

    # 校验数据
    if not args.force:
        errors = validate_data(data)
        if errors:
            print('❌ 数据校验失败（使用 --force 可跳过校验强制生成）:')
            for err in errors:
                print(f'   - {err}')
            sys.exit(1)
    else:
        print('⚠️ 已跳过数据校验（--force 模式）')

    try:
        output = build_docx(data, args.output)
    except Exception as e:
        print(f'❌ Word 文档生成失败: {e}')
        sys.exit(1)

    print(f'✅ Word 行程手册已生成: {output}')

    days = data.get('days', [])
    budget = data.get('budget') or {}
    print(f'   行程天数: {len(days)}')
    print(f'   目的地: {", ".join(data.get("destinations", []))}')
    print(f'   预估总费用: ¥{budget.get("grand_total", 0):,}')


if __name__ == '__main__':
    main()
